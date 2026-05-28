"""Generate Phase 2 report as a formatted .docx university document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

HERE = Path(__file__).parent

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    return p

def add_paragraph(doc, text="", bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    # light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # blue background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tc_pr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "FFFFFF" if r_idx % 2 == 0 else "EBF3FB"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=10)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table

def add_inline_label(p, label, value):
    """Add 'Label: value' within a paragraph."""
    r1 = p.add_run(label + ": ")
    set_font(r1, bold=True, size=11)
    r2 = p.add_run(value)
    set_font(r2, size=11)

def page_break(doc):
    doc.add_page_break()

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Visual Ingredient Scanner\nwith Recipe Generation")
set_font(r, size=26, bold=True, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Phase 2 Checkpoint Report")
set_font(r, size=16, italic=True, color=(0x44, 0x72, 0xC4))

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

info_lines = [
    ("Course",   "Computer Vision — Master MEI FIB · UPC · Spring 2026"),
    ("Team",     "Pol Plana · Emma Nájera · Houda El Fezzak"),
    ("Deadline", "26–28 May 2026"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_label(p, label, value)

doc.add_paragraph()
add_divider(doc)

page_break(doc)

# ── 1. PROJECT OVERVIEW ───────────────────────────────────────────────────────

add_heading(doc, "1. Project Overview", 1)

add_body(doc,
    "The Visual Ingredient Scanner is an end-to-end mobile application that allows a user "
    "to point a phone camera at a fridge or kitchen counter, tap once, and receive a list "
    "of detected ingredients with estimated weights and three ranked recipe suggestions."
)
add_body(doc,
    "The system is designed to be serverless and privacy-first: the computer vision "
    "pipeline runs entirely on-device, with only two lightweight API calls touching the "
    "network — one for ingredient density lookup (cached per class after the first call) "
    "and one for recipe generation."
)

add_heading(doc, "1.1  Five-Stage Pipeline", 2)

add_code_block(doc,
    "Camera frame\n"
    "     │\n"
    "     ▼\n"
    "① YOLO11s ──────────────► bounding boxes + class labels    (on-device)\n"
    "     │\n"
    "     ▼\n"
    "② Depth Anything V2-S ──► per-pixel depth map              (on-device)\n"
    "     │\n"
    "     ├── ③ Gemini density call ──► kg/m³ per class (cached) (cloud, once per new class)\n"
    "     │\n"
    "     ▼\n"
    "④ Pinhole model + shape heuristics ──► weight per item (g)  (on-device)\n"
    "     │\n"
    "     ▼\n"
    "⑤ Gemini recipe call ──────────────► 3 ranked recipes JSON  (cloud, once per scan)"
)

# ── 2. PHASE 2 STATUS ─────────────────────────────────────────────────────────

add_heading(doc, "2. Phase 2 Status Summary", 1)

add_table(doc,
    headers=["Deliverable", "Status"],
    rows=[
        ["Fine-tune YOLO11s on food dataset",     "✅  Done — mAP50-95 = 0.554"],
        ["Export Depth Anything V2-S to ONNX",    "✅  Done (models/depth/)"],
        ["Implement full 5-stage Python pipeline", "✅  Done (pipeline/)"],
        ["Build working Gradio laptop demo",       "✅  Done (prototype/app.py)"],
        ["Record detection metrics (mAP)",         "✅  Recorded — see §6"],
        ["Weight estimation functional",           "✅  Produces plausible gram estimates"],
        ["Phase 2 report",                         "✅  This document"],
    ],
    col_widths=[9, 8],
)

# ── 3. DATASET ────────────────────────────────────────────────────────────────

add_heading(doc, "3. Dataset", 1)

add_heading(doc, "3.1  Composition Strategy", 2)
add_body(doc,
    "No single public dataset covers all 68 target food classes with sufficient training "
    "instances. We merged four complementary Roboflow Universe datasets into a single "
    "Roboflow project, applying class renames to enforce consistent naming across all "
    "pipeline stages."
)

add_heading(doc, "3.2  Source Datasets", 2)
add_table(doc,
    headers=["Dataset", "Roboflow slug", "Main contribution"],
    rows=[
        ["Ingredient Detection",        "ingredient-detection-unorginazed-data", "Core fresh produce, proteins, dairy"],
        ["Veggies and Fruits Balanced", "veggies-and-fruits-balanced-0g1ss",     "Fruits, lettuce, exotic produce"],
        ["Vegetables Dataset",          "vegetables-g9p5a",                      "Zucchini, beet, cauliflower, celery"],
        ["Groceries",                   "groceries-mts9o",                       "Packaged goods: pasta, oil, cereal, juice"],
    ],
    col_widths=[4.5, 6, 6.5],
)

add_heading(doc, "3.3  Class Balance and Capping", 2)
add_body(doc,
    "After merging, dominant classes (e.g., orange: 11,085 instances) were capped at "
    "2,000 instances using an instance-count–based deletion strategy: images containing "
    "the most instances of the over-represented class were deleted first, preserving "
    "multi-class images where possible. This prevents class imbalance from biasing "
    "training towards the most frequent categories."
)

add_heading(doc, "3.4  Final Class List — 68 Classes", 2)

categories = [
    ("Fruits (22)", "apple, avocado, banana, blackberries, blueberries, cantaloupe, coconut, "
                    "fig, grapes, grapefruit, kiwi, lemon, lime, mango, orange, peach, pear, "
                    "pineapple, pomegranate, raspberries, strawberries, watermelon"),
    ("Vegetables (28)", "artichoke, beet, broccoli, brussels_sprouts, cabbage, carrot, cauliflower, "
                        "celery, chili, corn, cucumber, eggplant, garlic, ginger, green_beans, "
                        "lettuce, mushrooms, okra, onion, peas, pepper, potato, pumpkin, radish, "
                        "spinach, sweet_potato, tomato, zucchini"),
    ("Proteins & Dairy (12)", "beef, butter, cheese, chicken, egg, fish, ham, heavy_cream, pork, "
                               "shrimp, tofu, yogurt"),
    ("Pantry & Packaged (21)", "bread, cereal, chocolate, coffee, flour, honey, hummus, jam, juice, "
                                "mayonnaise, milk, nuts, oil, pasta, rice, soda, sugar, tea, "
                                "tomato_sauce, vinegar, water"),
]
for cat, items in categories:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(cat + ":  ")
    set_font(r1, bold=True, size=10.5)
    r2 = p.add_run(items)
    set_font(r2, size=10.5, italic=True)

doc.add_paragraph()

add_heading(doc, "3.5  Dataset Split", 2)
add_table(doc,
    headers=["Split", "Proportion", "Purpose"],
    rows=[
        ["Train",      "70 %", "YOLO11s fine-tuning"],
        ["Validation", "20 %", "Loss monitoring during training"],
        ["Test",       "10 %", "Final mAP evaluation (held-out)"],
    ],
    col_widths=[4, 4, 9],
)

# ── 4. MODEL CHOICES ─────────────────────────────────────────────────────────

page_break(doc)
add_heading(doc, "4. Model Choices and Architecture", 1)

add_heading(doc, "4.1  Stage ①  —  Object Detection: YOLO11s", 2)
add_body(doc,
    "YOLO11s (Ultralytics, small variant) was selected as the detection backbone. "
    "With approximately 9.4 M parameters and a 47.0 mAP50-95 baseline on COCO, it "
    "offers a strong accuracy-to-size trade-off suitable for on-device inference. "
    "YOLO11n was rejected for insufficient accuracy; YOLO11m was rejected as too large "
    "for mobile deployment. RF-DETR was excluded due to immature TFLite/CoreML export support."
)
add_table(doc,
    headers=["Property", "Value"],
    rows=[
        ["Architecture",      "Ultralytics YOLO11, small variant"],
        ["Parameters",        "~9.4 M"],
        ["Base training",     "COCO (80 classes)"],
        ["Fine-tuning",       "Food dataset, 68 classes, 25 epochs on Kaggle T4 GPU"],
        ["Android export",    "INT8 TFLite — ultralytics export format=tflite int8=True"],
        ["iOS export",        "CoreML — ultralytics export format=coreml"],
    ],
    col_widths=[5, 12],
)

add_heading(doc, "4.2  Stage ②  —  Depth Estimation: Depth Anything V2-S", 2)
add_body(doc,
    "Depth Anything V2-S (Apache 2.0) was chosen as the depth estimator. The small "
    "variant balances accuracy and inference speed for on-device use. V2-Base and "
    "V2-Large were excluded due to CC-BY-NC licence restrictions that prevent "
    "academic/commercial deployment. The model is used pretrained with no fine-tuning."
)
add_body(doc,
    "Depth calibration note: the current ONNX export uses the relative-depth checkpoint, "
    "which outputs unitless disparity values rather than calibrated metres. In the Python "
    "prototype, the depth map is normalised in two steps: (1) global min-max is mapped to "
    "the range [0.35 m, 3.0 m], representing a realistic indoor kitchen scene; "
    "(2) the map is then rescaled so the median depth across all detected food bounding "
    "boxes equals 0.50 m, anchoring to the typical phone-to-counter distance. "
    "A proper metric export (depth-anything/Depth-Anything-V2-Small-Metric-Indoor-hf) "
    "is planned before Phase 3."
)
add_table(doc,
    headers=["Property", "Value"],
    rows=[
        ["Checkpoint",   "depth-anything/Depth-Anything-V2-Small (Apache 2.0)"],
        ["Export",       "ONNX via torch.onnx.export, opset 17"],
        ["Runtime",      "onnxruntime (CPU)"],
        ["Fine-tuning",  "None — pretrained checkpoint used as-is"],
    ],
    col_widths=[5, 12],
)

add_heading(doc, "4.3  Stage ③  —  Density Lookup: Gemini API", 2)
add_body(doc,
    "Ingredient densities (kg/m³) are fetched from Gemini once per newly encountered "
    "class and cached in data/density_cache.json. All unseen classes from a single scan "
    "are batched into one API call. A static fallback table (data/density_fallback.json, "
    "68 entries) is used when the API is unavailable, ensuring the pipeline remains "
    "fully functional offline for weight estimation."
)

add_heading(doc, "4.4  Stage ④  —  Weight Estimation: Pinhole Model + Shape Heuristics", 2)
add_body(doc,
    "Real-world dimensions are derived from the detected bounding box and depth map "
    "using the standard pinhole camera model:"
)
add_code_block(doc,
    "real_width  = (bbox_width_px  / focal_length_px) × depth_m\n"
    "real_height = (bbox_height_px / focal_length_px) × depth_m\n"
    "\n"
    "weight_g = volume_m³ × density_kg_m³ × 1000"
)
add_body(doc,
    "Volume is estimated using per-class shape heuristics. Each class is assigned one "
    "of three geometric primitives:"
)
add_table(doc,
    headers=["Shape", "Classes (examples)", "Volume formula"],
    rows=[
        ["Sphere",   "apple, tomato, onion, egg, lemon", "V = (4/3)π(d/2)³,   d = min(w, h)"],
        ["Cylinder", "banana, carrot, cucumber, oil",    "V = π(w/2)² · h"],
        ["Box",      "bread, cheese, chicken, cereal",   "V = w · h · max(w, h) · 0.5"],
    ],
    col_widths=[3, 6.5, 7.5],
)
add_body(doc,
    "focal_length_px is read from the image EXIF tag FocalLengthIn35mmFilm and converted "
    "to pixels. When EXIF data is absent (e.g., screenshots, cropped images), the fallback "
    "is image_width × 0.8, which approximates a standard wide-angle phone lens."
)

add_heading(doc, "4.5  Stage ⑤  —  Recipe Generation: Gemini API", 2)
add_body(doc,
    "After weight estimation, detected ingredients and their estimated gram weights are "
    "sent to Gemini in a single structured prompt. The model returns a JSON array of "
    "three ranked recipes, each containing a name, ingredient list with quantities, "
    "ordered preparation steps, and serving count. Recipe quantities are adapted to "
    "the detected amounts (e.g., a single-serving dish if only 120 g of pasta is detected)."
)

# ── 5. IMPLEMENTATION ─────────────────────────────────────────────────────────

page_break(doc)
add_heading(doc, "5. Implementation", 1)

add_heading(doc, "5.1  Repository Structure", 2)
add_code_block(doc,
    "pipeline/\n"
    "├── detect.py       YOLO11s inference wrapper\n"
    "├── depth.py        Depth Anything V2-S ONNX wrapper + depth calibration\n"
    "├── density.py      Gemini density call + local JSON cache\n"
    "├── weight.py       Pinhole model + shape heuristics\n"
    "├── recipe.py       Gemini recipe generation\n"
    "└── pipeline.py     End-to-end orchestration\n"
    "\n"
    "training/\n"
    "├── train_yolo.py              YOLO11s fine-tuning script\n"
    "├── export_yolo.py             TFLite + CoreML export\n"
    "└── export_depth_onnx.py       Depth Anything V2-S → ONNX\n"
    "\n"
    "prototype/\n"
    "└── app.py          Gradio laptop demo\n"
    "\n"
    "data/\n"
    "├── classes.yaml               68 classes with shape hints and densities\n"
    "├── density_fallback.json      Static density table (68 entries)\n"
    "└── density_cache.json         Runtime Gemini density cache\n"
    "\n"
    "models/\n"
    "├── yolo/food_detector.pt      Fine-tuned YOLO11s (19.2 MB)\n"
    "└── depth/depth_anything_v2_small.onnx"
)

add_heading(doc, "5.2  Gradio Laptop Demo", 2)
add_body(doc,
    "A Gradio web interface (prototype/app.py) was built as the Phase 2 deliverable. "
    "The user uploads a kitchen photo, clicks Scan, and receives: (1) the input image "
    "annotated with bounding boxes and estimated weights, (2) a text list of detected "
    "ingredients and grams, and (3) three formatted recipe suggestions. All five pipeline "
    "stages run end-to-end on the laptop CPU."
)
add_code_block(doc,
    "# Activate virtual environment and launch\n"
    "venv\\Scripts\\activate\n"
    "python prototype/app.py\n"
    "# → open http://127.0.0.1:7860 in a browser"
)

add_heading(doc, "5.3  Training Setup", 2)
add_body(doc,
    "Training was executed on Kaggle with a T4 GPU using notebooks/train_yolo_kaggle.ipynb. "
    "Google Colab was initially used but abandoned due to unreliable Google Drive mounting "
    "and a training speed of approximately 20 minutes per epoch. Kaggle provided "
    "approximately 8 minutes per epoch and persistent output storage across kernel restarts."
)
add_body(doc, "Training configuration:")
for item in [
    "Model: YOLO11s pretrained on COCO",
    "Epochs: 25 with cosine learning rate schedule",
    "Batch size: 16, image size: 640 × 640",
    "Augmentations: built-in Ultralytics augmentations (mosaic, flip, scale, HSV jitter)",
    "Dataset: 68-class food dataset from Roboflow (70/20/10 split)",
    "Total training time: ~3.26 hours (11,754 s)",
]:
    add_bullet(doc, item)

doc.add_paragraph()

# ── 6. RESULTS ────────────────────────────────────────────────────────────────

page_break(doc)
add_heading(doc, "6. Results", 1)

add_heading(doc, "6.1  Detection — YOLO11s mAP", 2)
add_body(doc,
    "The fine-tuned YOLO11s was evaluated on the held-out test split (10% of the merged "
    "dataset). The model comfortably exceeds the Phase 2 target of mAP50-95 > 0.40."
)
add_table(doc,
    headers=["Metric", "Value", "Target", "Status"],
    rows=[
        ["mAP50-95 (primary)", "0.554",  "> 0.40", "✅  Exceeded"],
        ["mAP50",              "0.744",  "—",       "—"],
        ["Precision",          "0.745",  "—",       "—"],
        ["Recall",             "0.699",  "—",       "—"],
        ["Training epochs",    "25",     "—",       "—"],
        ["Training time",      "3.26 h", "—",       "—"],
    ],
    col_widths=[5.5, 3.5, 3.5, 4.5],
)

add_body(doc,
    "Training loss curves show consistent convergence over 25 epochs with no signs of "
    "overfitting: both train and validation box loss decrease steadily, and the mAP50-95 "
    "curve crosses the 0.40 target at approximately epoch 9 and continues to improve "
    "through epoch 25."
)

add_body(doc, "Selected per-class performance highlights:")
add_table(doc,
    headers=["Class", "Approx. mAP50-95", "Notes"],
    rows=[
        ["ginger",       "0.881", "Very distinctive appearance, high accuracy"],
        ["blackberries", "0.841", "Unique texture — easiest class to detect"],
        ["jam",          "0.834", "Consistent label design in training data"],
        ["garlic",       "0.396", "Visually similar to onion — moderate performance"],
        ["orange",       "0.065", "Few validation images after instance capping"],
        ["pomegranate",  "0.082", "Small validation set — not representative"],
    ],
    col_widths=[4.5, 4, 8.5],
)

add_heading(doc, "6.2  Weight Estimation", 2)
add_body(doc,
    "Weight estimation is functional and produces plausible gram estimates for close-up "
    "kitchen photography at approximately 50 cm phone-to-food distance. Absolute accuracy "
    "is currently limited by the absence of a calibrated metric depth model. "
    "The following results were obtained on a test image containing five detected items:"
)
add_table(doc,
    headers=["Item", "Estimated weight (g)", "Typical real weight (g)", "Approx. error"],
    rows=[
        ["Lemon",       "150–250",  "~100",     "×1.5–2.5"],
        ["Onion",       "150–200",  "~150",     "×1.0–1.3"],
        ["Pomegranate", "350–500",  "~300",     "×1.2–1.7"],
        ["Tomato",      "400–700",  "150–250",  "×2.0–4.0"],
        ["Apple",       "500–900",  "180–250",  "×2.5–4.5"],
    ],
    col_widths=[3.5, 4.5, 4.5, 4.5],
)
add_body(doc,
    "Rounder, more compact items (lemon, onion, pomegranate) are estimated within a "
    "factor of 1.5–2× of the real weight. Items with looser bounding boxes relative to "
    "their actual footprint (tomato, apple) show higher relative error. Replacing the "
    "heuristic depth calibration with a metric ONNX export is expected to reduce these "
    "errors significantly."
)

add_heading(doc, "6.3  Depth Estimation", 2)
add_body(doc,
    "A quantitative δ₁ accuracy evaluation (percentage of pixels within 25% of ground "
    "truth depth) was not performed at this stage, as it requires a paired RGB + LiDAR "
    "ground-truth dataset not available in our setup. Qualitative inspection of the "
    "estimated depth maps shows correct relative ordering — closer objects are assigned "
    "lower depth values and farther objects higher values — with plausible spatial "
    "structure across kitchen scenes."
)

# ── 7. KNOWN LIMITATIONS ──────────────────────────────────────────────────────

page_break(doc)
add_heading(doc, "7. Known Limitations", 1)

limitations = [
    ("Depth model not metric",
     "The current ONNX export uses the relative-depth checkpoint of Depth Anything V2-S. "
     "Depth values are calibrated via a two-step heuristic: global scene normalisation "
     "to [0.35 m, 3.0 m] followed by anchoring the median food depth to 0.50 m. "
     "A proper metric export is planned before Phase 3."),
    ("Weight accuracy varies by item shape",
     "Round, compact items are estimated within a factor of 1.5–2× of actual weight. "
     "Items with irregular shapes or loose bounding boxes may be overestimated by 3–5×. "
     "The ±30% MAPE target requires both a metric depth model and tighter bounding box crops."),
    ("Packaged goods detection is weak",
     "Classes such as pasta, oil, and juice are rarely detected with sufficient confidence. "
     "Packaging appearance varies widely across brands; additional training data or a "
     "dedicated packaging-aware detection head would be required to improve these classes."),
    ("Gemini API free-tier quota",
     "The gemini-2.0-flash-lite and gemini-2.0-flash models showed a hard limit of 0 "
     "requests on the free tier in certain Google account configurations. The pipeline "
     "was migrated to the google-genai SDK (≥1.0) and gemini-1.5-flash. "
     "Density estimation already works fully offline via the static density_fallback.json."),
    ("Low-instance classes",
     "mayonnaise (42 training instances) and hummus (109 instances) are below the "
     "recommended threshold of 150+ instances per class. Detection recall for these "
     "classes is below average; additional data can be added before Phase 3."),
]

for i, (title, body) in enumerate(limitations, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{i}.  {title}")
    set_font(r, bold=True, size=11)
    add_body(doc, body)

# ── 8. NEXT STEPS ─────────────────────────────────────────────────────────────

add_heading(doc, "8. Next Steps — Phase 3 (due 15–17 June 2026)", 1)

next_steps = [
    "Re-export Depth Anything V2-S from the metric indoor checkpoint and remove the heuristic depth calibration layers.",
    "Run evaluation/eval_weight.py with the metric model to obtain a quantitative MAPE measurement on held-out items.",
    "Build the Flutter mobile application — all screens (scan, result) and services (detector, depth, weight, density, recipe) in mobile/.",
    "Validate the Depth Anything V2-S ONNX model on Android via onnxruntime_flutter; apply INT8 quantisation if operator compatibility issues arise.",
    "Export the fine-tuned YOLO11s to INT8 TFLite and integrate into the Flutter app via tflite_flutter.",
    "Measure end-to-end latency on a physical Android phone (target: < 5 s from capture to displayed results).",
    "Produce the final project report, presentation slides, and recorded demo video.",
]

for step in next_steps:
    add_bullet(doc, step)

doc.add_paragraph()

# ── SAVE ──────────────────────────────────────────────────────────────────────

out_path = HERE / "phase2_report.docx"
doc.save(str(out_path))
print(f"Report saved to {out_path}")
